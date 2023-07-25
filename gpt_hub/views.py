from django.contrib.auth import authenticate
from django.conf import settings
from django.middleware import csrf
from django.http import JsonResponse
from rest_framework import exceptions as rest_exceptions, response, decorators as rest_decorators, permissions as rest_permissions
from rest_framework_simplejwt import tokens, views as jwt_views, serializers as jwt_serializers, exceptions as jwt_exceptions
from user import serializers, models
import traceback
import uuid
import os
import pinecone
from decouple import config
from PyPDF2 import PdfReader
from docx import Document
import openai

apiKey = config('OPENAI_API_KEY')
openai.api_key = apiKey

pinecone.init(api_key=config('PINECONE_API_KEY'),
                    environment=config('PINECONE_ENV'))
# Create your views here.

@rest_decorators.api_view(["POST"])
@rest_decorators.permission_classes([rest_permissions.IsAuthenticated])
def document_create(request):
    try:
        user_id  = request.user.id
        file = request.FILES['file']
        topic = request.POST['topic']
        time = request.POST['time']
        doc_type = request.POST['type']
        unique_filename = str(uuid.uuid4()) + "_" + file.name
        docs_directory = 'docs'
        file_path = os.path.join(settings.MEDIA_ROOT, docs_directory, unique_filename)
        content = []

        try:
            if doc_type == "PDF":
                content, flag = parse_pdf(file)
                if not flag:
                    return response.Response(({'message': 'Parsing Error or Empty PDF, Please Use other PDF'}))
            elif doc_type == "DOC":
                content, flag = parse_docx(file)
                if not flag:
                    return response.Response(({'message': 'Parsing Error or Empty Doc, Please use other Doc'}))
            else:
                return response.Response(({'message': 'Non-allowed Doc Types'}))
        except:
            return response.Response(({'message': 'Non-alled Doc Types'}))

        sentences = []
        chunk = ""
        # sume 5 sentences as a one chunk
        for i in range(0, len(content)):
            chunk += f"\n {content[i]}"
            if (i+1)%5 == 0:
                sentences.append(chunk)
                chunk = ""
        sentences_list , embeddings, vec_indexs = get_embedding(sentences)
        if len(embeddings) == 0:
            return response.Response(({"message": "Creating Embedding Error"}))
        
        meta =[{"sentence": line} for line in sentences_list]

        vector = list(zip(vec_indexs, embeddings, meta))

        isCreatingEmbedding = embedding_to_pinecone(vector, topic)

        if not isCreatingEmbedding:
                return response.Response(({'message': "Inserting Embeeding Error"}))

        new_document = models.Document(
            user_id=user_id,
            document_type=doc_type,
            date=time,
            topic=topic,
            file=file_path
        )
        new_document.save()

        with open(file_path, 'wb') as destination:
            for chunk in file.chunks():
                destination.write(chunk)
        return response.Response(({'message': True}))
    
    except Exception as e:
        print(traceback.format_exc())
        return response.Response(({'message': False}))

@rest_decorators.api_view(["POST"])
@rest_decorators.permission_classes([rest_permissions.IsAuthenticated])
def chat_doc(request):
    query = request.data['query']
    topic = request.data['topic']
    queryResponse = query_embedding(query, topic)
    if not queryResponse:
            return response.Response({"message": "Querying Embedding Error"})
    inputSentence = ""

    for i in queryResponse["matches"]:
            inputSentence += i["metadata"]["sentence"]
    inputSentence = limit_string_tokens(inputSentence, 2500)

    try:
        res = openai.ChatCompletion.create(
            model = "gpt-3.5-turbo",
            temperature = 0.1,
            messages = [
                {"role": "system", "content" : "You are a helpful assistant."},
                {"role": "user", "content": inputSentence},
                {"role": "user", "content": query},
            ]
        )
        result = res["choices"][0]["message"]["content"]
        return response.Response({"message": result})
    except Exception as e:
        print(traceback.format_exc())
        return response.Response({"message": "Net Error"})

def limit_string_tokens(string, max_tokens):
    tokens = string.split()  # Split the string into tokens
    if len(tokens) <= max_tokens:
        return string  # Return the original string if it has fewer or equal tokens than the limit
    limited_string = ' '.join(tokens[:max_tokens])
    return limited_string

def query_embedding(question, nameSpace):
    sentences, embeddings, vec_indexs = get_embedding([question])
    if len(embeddings) == 0:
        return False
    try:
        pinecone.init(api_key=config('PINECONE_API_KEY'),
                    environment=config('PINECONE_ENV'))
        active_indexes = pinecone.list_indexes()
        index = pinecone.Index(active_indexes[0])
        query_response = index.query(
                namespace=nameSpace,
                top_k=50,
                include_values=True,
                include_metadata=True,
                vector=embeddings[0],
            )
        
        return query_response
    except Exception as e:
        print(traceback.format_exc())
        return False

def get_embedding(content):
    try:
        # Embed a line of text
        response = openai.Embedding.create(
            model= "text-embedding-ada-002",
            input=content
        )
        embedding = []
        vec_indexs = []
        # Extract the AI output embedding as a list of floats
        # embedding = response["data"][0]["embedding"]
        index = 0
        for i in response["data"]:
            index += 1
            embedding.append(i["embedding"])
            vec_indexs.append("vec" + str(index))
        # creating the vector indexes
        return content, embedding, vec_indexs
    except Exception as e:
        print(traceback.format_exc())
        return [], [], []

def embedding_to_pinecone(vector, nameSpace):
    # Initialized pinecone client
    try:
        # Testing the indexs client
        active_indexes = pinecone.list_indexes()
        if len(active_indexes) != 0:
            index = pinecone.Index(active_indexes[0])
            print(index)
            try:
                # inserting the embedding
                vectors_list = chunk_list(vector, 50)
                for i in vectors_list:
                    print(i)
                    index.upsert(vectors=i, namespace=nameSpace)
                print("Successfull inserted embeddings")
            except Exception as e:
                print("Error inserting embeddings:")
                print(traceback.format_exc())
        else:
            print("create index")
            pinecone.create_index("example-index", dimension=1536)
        return True
    except Exception as e:
        print(traceback.format_exc())
        return False

def chunk_list(input_list, chunk_size):
    return [input_list[i:i + chunk_size] for i in range(0, len(input_list), chunk_size)]

def parse_pdf(uploadedFile):
    try:
        pdf = PdfReader(uploadedFile)
        parsedData = {
            'num_pages': len(pdf.pages),
            'text': [page.extract_text() for page in pdf.pages]
        }
        content = []
        for pageData in parsedData['text']:
            for each_sentence in pageData.split("\n"):
                if len(each_sentence) > 2:
                    content.append(each_sentence)
        flag =  True
        if len(content) == 0:
            flag = False
        return content, flag
    except Exception as e:
        print(traceback.format_exc())
        return [] , flag
    
def parse_docx(uploadedFile):
    try:
        doc = Document(uploadedFile)
        content = []
        for paragraph in doc.paragraphs:
            for sentence in paragraph.text.split("\n"):
                if len(sentence) > 2:
                    content.append(sentence)
        flag = True
        if len(content) == 0:
            flag = False
        return content, flag
    except Exception as e:
        print(traceback.format_exc())
        return [], flag
